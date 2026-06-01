"""Generate supervisor-ready deliverables from the Week 1-2 work:

  1. Progress_Report_Week1-2.docx  (Word document)
  2. Supervisor_Meeting_Week1-2.pdf (slide deck PDF with embedded figures)

Run inside the project venv:
    python scripts/make_deliverables.py
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image as PILImage

# ----- docx -----
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----- reportlab (PDF slides) -----
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, Image, PageBreak)

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "outputs" / "figures"
HULL_BLUE = RGBColor(0x1F, 0x3A, 0x6E)
PDF_BLUE = colors.HexColor("#1F3A6E")
PDF_GREEN = colors.HexColor("#1B7F4B")


# =====================================================================
# 1. PROGRESS REPORT  ->  DOCX
# =====================================================================
def build_docx(out_path: Path) -> None:
    doc = Document()

    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = HULL_BLUE
        return h

    def table_from_rows(rows, header=True):
        t = doc.add_table(rows=0, cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
                if header and i == 0:
                    for p in cells[j].paragraphs:
                        for r in p.runs:
                            r.font.bold = True
        return t

    # --- Title block ---
    title = doc.add_heading("Progress Report — Weeks 1–2 (Foundation Phase)", level=0)
    for r in title.runs:
        r.font.color.rgb = HULL_BLUE

    meta = [
        ("Project", "Classification of Malignant Melanoma in Canines (CATCH dataset)"),
        ("Project Code", "DAIM2025A_088"),
        ("Student", "Muhammad Tayyab Abbas"),
        ("Supervisor", "Dr Claire Cashmore"),
        ("Reporting period", "Week 1–2 (26 May – 8 June 2026)"),
        ("Meeting", "First fortnightly supervisor meeting"),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.style = "Light List Accent 1"
    for k, v in meta:
        c = mt.add_row().cells
        c[0].text = k
        c[1].text = v
        for p in c[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
    doc.add_paragraph()

    # --- 1. Summary ---
    heading("1. Summary", 1)
    doc.add_paragraph(
        "The Foundation phase is on schedule. The computational environment is fully "
        "configured inside an isolated virtual environment, the literature review has "
        "been refined, and the data acquisition and preprocessing pipeline (Macenko "
        "stain normalisation → patch extraction → stratified split) has been "
        "implemented and validated end-to-end. The pipeline was verified on synthetic "
        "H&E slides while the CATCH dataset download from TCIA is being finalised; it is "
        "ready to run unchanged on the real slides."
    )

    # --- 2. Planned vs actual ---
    heading("2. Planned vs. actual (Gantt chart)", 1)
    table_from_rows([
        ["Gantt task", "Planned weeks", "Status", "Evidence"],
        ["Environment Setup & Framework", "W1", "Complete",
         ".venv, requirements.txt, requirements-lock.txt, setup_env scripts"],
        ["Literature Review & Refinement", "W1–W2", "Complete",
         "docs/literature/literature_review_notes.md"],
        ["Data Acquisition & Preprocessing", "W1–W3", "In progress (~60%)",
         "code + QA/split reports + figures"],
    ])

    # --- 3. What was completed ---
    heading("3. What was completed", 1)

    heading("3.1 Environment setup (Week 1)", 2)
    for b in [
        "Python 3.12 in an ISOLATED virtual environment (.venv) so project packages "
        "never affect other projects or the global interpreter.",
        "Dependencies: OpenCV, scikit-image, scikit-learn, OpenSlide, NumPy, pandas, "
        "matplotlib (PyTorch + Albumentations added for Week 3 model development).",
        "Reproducible setup: requirements.txt, environment.yml, pinned "
        "requirements-lock.txt, and one-command installers (setup_env.ps1 / .sh).",
        "Git repository with a clean, modular src/ package structure.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    heading("3.2 Literature review refinement (Weeks 1–2)", 2)
    for b in [
        "Consolidated the key references underpinning the methodology (U-Net, "
        "Attention U-Net, ResNet, EfficientNet, Macenko, Grad-CAM, canine "
        "comparative-oncology papers).",
        "Method and finding summaries recorded in "
        "docs/literature/literature_review_notes.md.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    heading("3.3 Data acquisition and preprocessing (Weeks 1–3, in progress)", 2)
    doc.add_paragraph("Implemented and validated the complete preprocessing chain:")
    for i, b in enumerate([
        "Acquisition helper (download_catch.py) — documents the TCIA / NBIA Data "
        "Retriever workflow and verifies downloaded slides.",
        "Quality assessment (quality_check.py) — per-slide checks for size, blur, "
        "brightness and tissue fraction; flags/excludes poor slides with a reason.",
        "Macenko stain normalisation (stain_normalization.py) — maps slides onto a "
        "common H&E colour appearance to remove scanner/lab variation.",
        "Patch extraction (patch_extraction.py) — tiles slides into 256×256 patches, "
        "keeping only tissue-bearing tiles (≥50% tissue).",
        "Stratified split (dataset_split.py) — 70/15/15 train/val/test split "
        "preserving class distribution.",
    ], 1):
        doc.add_paragraph(f"{i}. {b}", style="List Number")
    p = doc.add_paragraph()
    p.add_run("Validation run (synthetic slides): ").bold = True
    p.add_run("2 slides → both passed QA → 26 tissue patches extracted → split "
              "18/4/4. Confirms the pipeline runs end-to-end and is ready for the "
              "real CATCH slides.")

    # --- 4. Artefacts ---
    heading("4. Artefacts produced (shown in this meeting)", 1)
    table_from_rows([
        ["Artefact", "Location"],
        ["Before/after stain normalisation figures", "outputs/figures/stain_normalization_*.png"],
        ["Sample extracted patches (grids)", "outputs/figures/patches_*.png"],
        ["Quality assessment report", "outputs/logs/quality_report.csv"],
        ["Train/val/test split summary", "outputs/logs/split_summary.json"],
        ["Source code (modular package)", "src/"],
    ])

    # embed the key figure if present
    fig = FIG / "stain_normalization_demo_slide_bluish.png"
    if fig.exists():
        doc.add_paragraph()
        cap = doc.add_paragraph("Figure: Macenko stain normalisation (before / after).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(fig), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 5. Risks ---
    heading("5. Risks & issues (current)", 1)
    table_from_rows([
        ["Risk", "Status", "Mitigation in place"],
        ["CATCH download (TCIA account / large files)", "Active",
         "Pipeline validated on synthetic data; runs unchanged on real slides once "
         "downloaded. NBIA Data Retriever workflow documented."],
        ["GPU resources for upcoming training", "Watching",
         "Google Colab Pro + university HPC identified as fallback."],
        ["Class imbalance across tumour subtypes", "Anticipated",
         "Stratified split implemented; class-weighted loss + augmentation planned for Week 3+."],
    ])

    # --- 6. Plan ---
    heading("6. Plan for next fortnight (Weeks 3–5)", 1)
    for b in [
        "Complete CATCH download and run the validated pipeline on all real slides.",
        "Finalise patch dataset and class distribution statistics.",
        "Begin U-Net segmentation model (baseline → ResNet-34 encoder), with "
        "TensorBoard monitoring of loss convergence.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # --- 7. Questions ---
    heading("7. Questions for supervisor", 1)
    for b in [
        "Is the QA exclusion threshold (≥10% tissue, blur ≥ 50) appropriate, or "
        "should it be agreed with a veterinary pathologist?",
        "Confirm the final list of tumour subtypes to include as classification classes.",
        "Any preference on patch magnification priority (5×/10×/20×) for the first "
        "segmentation experiments?",
    ]:
        doc.add_paragraph(b, style="List Number")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[DOCX] saved -> {out_path}")


# =====================================================================
# 2. SLIDE DECK  ->  PDF (with embedded figures)
# =====================================================================
def _img_flowable(path: Path, max_w: float, max_h: float):
    """Return a reportlab Image scaled to fit within (max_w, max_h)."""
    with PILImage.open(path) as im:
        iw, ih = im.size
    ratio = min(max_w / iw, max_h / ih)
    return Image(str(path), width=iw * ratio, height=ih * ratio)


def build_pdf(out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    page = landscape(A4)
    doc = SimpleDocTemplate(str(out_path), pagesize=page,
                            leftMargin=42, rightMargin=42,
                            topMargin=40, bottomMargin=40)
    usable_w = page[0] - 84
    usable_h = page[1] - 80

    ss = getSampleStyleSheet()
    h_title = ParagraphStyle("t", parent=ss["Title"], textColor=PDF_BLUE,
                             fontSize=30, leading=36, spaceAfter=10)
    h_sub = ParagraphStyle("s", parent=ss["Normal"], fontSize=15, leading=20,
                           textColor=colors.HexColor("#333333"))
    h_slide = ParagraphStyle("sh", parent=ss["Heading1"], textColor=PDF_BLUE,
                             fontSize=22, leading=26, spaceAfter=14)
    body = ParagraphStyle("b", parent=ss["Normal"], fontSize=14, leading=22)
    bullet = ParagraphStyle("bl", parent=body, leftIndent=16, bulletIndent=4,
                            spaceAfter=6)
    caption = ParagraphStyle("cap", parent=ss["Italic"], fontSize=11,
                             alignment=1, textColor=colors.HexColor("#555555"))

    story = []

    def slide(flowables):
        story.extend(flowables)
        story.append(PageBreak())

    def bullets(items):
        return [Paragraph(f"• {x}", bullet) for x in items]

    def make_table(data, col_widths=None, header_bg=PDF_BLUE):
        t = Table(data, colWidths=col_widths, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 13),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#EEF2F8")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ]))
        return t

    # Slide 1 — title
    slide([
        Spacer(1, 90),
        Paragraph("Classification of Malignant Melanoma in Canines", h_title),
        Paragraph("Week 1–2 Progress — Foundation Phase", h_sub),
        Spacer(1, 22),
        Paragraph("Muhammad Tayyab Abbas&nbsp;&nbsp;·&nbsp;&nbsp;MSc AI &amp; Data "
                  "Science&nbsp;&nbsp;·&nbsp;&nbsp;University of Hull", h_sub),
        Paragraph("Supervisor: Dr Claire Cashmore&nbsp;&nbsp;·&nbsp;&nbsp;"
                  "1st fortnightly meeting", h_sub),
    ])

    # Slide 2 — where we are
    slide([
        Paragraph("Where we are in the plan", h_slide),
        make_table([
            ["Phase", "Weeks", "Now"],
            ["Foundation", "1–3", "← we are here (W2)"],
            ["Model Development", "3–8", ""],
            ["Evaluation & Analysis", "8–10", ""],
            ["Documentation", "8–12", ""],
        ], col_widths=[260, 120, 240]),
        Spacer(1, 18),
        Paragraph("12-week project: 26 May → 14 Aug 2026, submission 17 Aug 2026.", body),
    ])

    # Slide 3 — planned vs done
    slide([
        Paragraph("Planned vs. done (Gantt)", h_slide),
        make_table([
            ["Task", "Weeks", "Status"],
            ["Environment Setup & Framework", "W1", "Done"],
            ["Literature Review & Refinement", "W1–2", "Done"],
            ["Data Acquisition & Preprocessing", "W1–3", "~60%"],
        ], col_widths=[420, 120, 160]),
        Spacer(1, 18),
        Paragraph("On schedule.", h_sub),
    ])

    # Slide 4 — environment
    slide([
        Paragraph("1 · Environment setup ✓", h_slide),
        *bullets([
            "Isolated virtual environment (.venv) — other projects untouched",
            "Python 3.12 + OpenCV, scikit-image, scikit-learn, OpenSlide (PyTorch next)",
            "Reproducible: requirements.txt + requirements-lock.txt + setup_env scripts",
            "Clean modular src/ package, version-controlled (Git)",
        ]),
    ])

    # Slide 5 — literature
    slide([
        Paragraph("2 · Literature review refined ✓", h_slide),
        *bullets([
            "Segmentation: U-Net (Ronneberger 2015), Attention U-Net (Oktay 2018)",
            "Classification: ResNet-50 (He 2016), EfficientNet-B3 (Tan & Le 2019)",
            "Preprocessing: Macenko stain normalisation (2009)",
            "Interpretability: Grad-CAM (Selvaraju 2017)",
            "Comparative oncology: Gillard (2014), Prouteau & André (2019)",
        ]),
    ])

    # Slide 6 — pipeline
    slide([
        Paragraph("3 · Preprocessing pipeline — built & validated ✓", h_slide),
        Paragraph("Whole-slide image → QA → Macenko normalise → patch extraction "
                  "→ 70/15/15 split", body),
        Spacer(1, 12),
        *bullets([
            "Quality check flags blurry / low-tissue / washed-out slides automatically",
            "Macenko removes scanner/lab colour variation",
            "Only tissue patches (≥50%) are kept",
            "Stratified split preserves class balance",
        ]),
    ])

    # Slide 7 — stain normalisation image
    fig1 = FIG / "stain_normalization_demo_slide_bluish.png"
    s7 = [Paragraph("Result — stain normalisation (before / after)", h_slide)]
    if fig1.exists():
        s7.append(_img_flowable(fig1, usable_w, usable_h - 130))
        s7.append(Spacer(1, 6))
        s7.append(Paragraph("Bluish slide mapped onto the standard H&E colour "
                            "appearance.", caption))
    slide(s7)

    # Slide 8 — patches image
    fig2 = FIG / "patches_demo_slide_pinkish.png"
    s8 = [Paragraph("Result — extracted tissue patches", h_slide)]
    if fig2.exists():
        s8.append(_img_flowable(fig2, usable_w, usable_h - 130))
        s8.append(Spacer(1, 6))
        s8.append(Paragraph("256×256 patches; background tiles automatically "
                            "discarded.", caption))
    slide(s8)

    # Slide 9 — validation run
    slide([
        Paragraph("Validation run (synthetic slides)", h_slide),
        make_table([
            ["Metric", "Value"],
            ["Slides processed", "2"],
            ["Passed QA", "2 / 2"],
            ["Patches extracted", "26"],
            ["Split (train/val/test)", "18 / 4 / 4"],
        ], col_widths=[360, 200]),
        Spacer(1, 16),
        Paragraph("Pipeline runs end-to-end and is ready to run unchanged on the "
                  "real CATCH slides once the TCIA download completes.", h_sub),
    ])

    # Slide 10 — risks
    slide([
        Paragraph("Risks & mitigations", h_slide),
        make_table([
            ["Risk", "Mitigation"],
            ["TCIA download / large files",
             "Pipeline validated on synthetic data; NBIA workflow documented"],
            ["GPU for training", "Colab Pro + university HPC fallback"],
            ["Class imbalance", "Stratified split now; weighted loss + augmentation next"],
        ], col_widths=[260, 460]),
    ])

    # Slide 11 — next
    slide([
        Paragraph("Next fortnight (Weeks 3–5)", h_slide),
        *bullets([
            "Complete CATCH download → run pipeline on all real slides",
            "Finalise patch dataset + class statistics",
            "Start U-Net segmentation (baseline → ResNet-34 encoder)",
        ]),
        Spacer(1, 14),
        Paragraph("Questions: QA thresholds OK? · final class list? · magnification "
                  "priority?", h_sub),
    ])

    # Slide 12 — thank you
    story.extend([
        Spacer(1, 150),
        Paragraph("Thank you", h_title),
        Paragraph("Questions &amp; feedback welcome", h_sub),
    ])

    doc.build(story)
    print(f"[PDF]  saved -> {out_path}")


def main():
    build_docx(ROOT / "docs" / "progress_report" / "Progress_Report_Week1-2.docx")
    build_pdf(ROOT / "docs" / "slides" / "Supervisor_Meeting_Week1-2.pdf")


if __name__ == "__main__":
    main()
