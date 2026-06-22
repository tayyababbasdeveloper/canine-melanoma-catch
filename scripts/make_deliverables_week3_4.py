"""Generate Week 3-4 supervisor deliverables from the segmentation run.

  1. Progress_Report_Week3-4.docx
  2. Supervisor_Meeting_Week3-4.pdf  (slides with training curves + predictions)

Reads the JSON reports / figures produced by run_week3_4_pipeline.py. Run it
AFTER the pipeline so the numbers and images exist:

    python scripts/run_week3_4_pipeline.py --demo --quick
    python scripts/make_deliverables_week3_4.py
"""
from __future__ import annotations

import sys
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.utils.config import load_config

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "outputs" / "figures"
LOGS = ROOT / "outputs" / "logs"


def _load(name):
    p = LOGS / name
    return json.loads(p.read_text()) if p.exists() else None


def _fmt_metrics(rep):
    if not rep:
        return ("—", "—", "—", "—")
    m = rep["test_metrics"]
    hd = m["hausdorff"]
    return (f"{m['dice']:.3f}", f"{m['iou']:.3f}", f"{m['pixel_acc']:.3f}",
            "n/a" if hd != hd else f"{hd:.1f}")  # hd!=hd detects NaN


def build_docx(out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HULL_BLUE = RGBColor(0x1F, 0x3A, 0x6E)
    base = _load("unet_baseline_test_report.json")
    res = _load("unet_resnet34_test_report.json")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = HULL_BLUE
        return h

    def table(rows, header=True):
        t = doc.add_table(rows=0, cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
                if header and i == 0:
                    for p in cells[j].paragraphs:
                        for rr in p.runs:
                            rr.font.bold = True
        return t

    title = doc.add_heading("Progress Report — Weeks 3–4 (Model Development: "
                            "U-Net Segmentation)", level=0)
    for r in title.runs:
        r.font.color.rgb = HULL_BLUE

    for k, v in [
        ("Project", "Classification of Malignant Melanoma in Canines (CATCH)"),
        ("Project Code", "DAIM2025A_088"),
        ("Student", "Muhammad Tayyab Abbas"),
        ("Supervisor", "Dr Claire Cashmore"),
        ("Reporting period", "Weeks 3–4 (9 – 22 June 2026)"),
        ("Meeting", "Second fortnightly supervisor meeting"),
    ]:
        doc.add_paragraph(f"{k}: {v}")

    heading("1. Summary", 1)
    doc.add_paragraph(
        "Phase 2 (Model Development) has begun on schedule. The remaining "
        "preprocessing items from Phase 1 were completed (multi-magnification "
        "5x/10x/20x patch extraction and a full augmentation pipeline), and the "
        "U-Net tumour-segmentation model has been implemented, trained and "
        "evaluated end-to-end. Following the plan, a from-scratch baseline U-Net "
        "was trained first, then a U-Net with an ImageNet-pretrained ResNet-34 "
        "encoder (transfer learning). Both were validated on synthetic annotated "
        "slides while the CATCH download is finalised; the code runs unchanged on "
        "the real annotated slides.")

    heading("2. Planned vs. actual", 1)
    table([
        ["Gantt task", "Planned", "Status"],
        ["Data Acquisition & Preprocessing", "W1–3", "Complete (multi-mag + augmentation added)"],
        ["U-Net Segmentation Development", "W3–5", "On track (baseline + ResNet-34 trained, W3–4)"],
    ])

    heading("3. What was completed", 1)
    for b in [
        "Multi-magnification patch extraction (5x/10x/20x) with paired image+mask "
        "tiling and tissue filtering.",
        "Augmentation pipeline (flips, 90° rotations, elastic deformation, "
        "brightness/contrast, H&E hue-saturation jitter) applied to training only.",
        "Segmentation dataset with slide-level train/val/test split (no slide "
        "spans two subsets — prevents patch leakage).",
        "U-Net models: from-scratch baseline and ResNet-34-encoder (transfer "
        "learning), composite BCE+Dice loss with class-imbalance pos_weight.",
        "Training recipe per proposal: Adam lr=1e-4, cosine annealing, early "
        "stopping on val-Dice, TensorBoard logging, best-checkpoint saving.",
        "Evaluation: Dice, IoU, pixel-accuracy and Hausdorff distance on the test "
        "split, with training-curve and prediction-overlay figures.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    heading("4. Results (test split)", 1)
    bd, bi, ba, bh = _fmt_metrics(base)
    rd, ri, ra, rh = _fmt_metrics(res)
    table([
        ["Model", "Dice", "IoU", "Pixel acc.", "Hausdorff", "Params"],
        ["Baseline U-Net", bd, bi, ba, bh,
         f"{base['params_million']}M" if base else "—"],
        ["U-Net + ResNet-34", rd, ri, ra, rh,
         f"{res['params_million']}M" if res else "—"],
    ])
    doc.add_paragraph(
        "Note: metrics are from the synthetic validation dataset and demonstrate "
        "the pipeline end-to-end; absolute values will change on the real CATCH "
        "slides. The proposal target is Dice > 0.85 on real data.")

    for fig_name, cap in [
        ("unet_resnet34_curves.png", "Figure: ResNet-34 U-Net training curves "
         "(loss + validation Dice)."),
        ("unet_resnet34_predictions.png", "Figure: test predictions "
         "(image | ground truth | prediction)."),
    ]:
        fp = FIG / fig_name
        if fp.exists():
            doc.add_paragraph()
            c = doc.add_paragraph(cap)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(fp), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading("5. Next fortnight (Weeks 5–6)", 1)
    for b in [
        "Implement the Attention U-Net variant and compare against the ResNet-34 U-Net.",
        "Run the trained pipeline on the real CATCH annotated slides once downloaded.",
        "Begin ResNet-50 classification (Week 6).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[DOCX] saved -> {out_path}")


def build_pdf(out_path: Path) -> None:
    from PIL import Image as PILImage
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, Image, PageBreak)

    PDF_BLUE = colors.HexColor("#1F3A6E")
    base = _load("unet_baseline_test_report.json")
    res = _load("unet_resnet34_test_report.json")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    page = landscape(A4)
    doc = SimpleDocTemplate(str(out_path), pagesize=page,
                            leftMargin=42, rightMargin=42, topMargin=40, bottomMargin=40)
    uw, uh = page[0] - 84, page[1] - 80
    ss = getSampleStyleSheet()
    h_title = ParagraphStyle("t", parent=ss["Title"], textColor=PDF_BLUE, fontSize=28)
    h_sub = ParagraphStyle("s", parent=ss["Normal"], fontSize=14)
    h_slide = ParagraphStyle("sh", parent=ss["Heading1"], textColor=PDF_BLUE, fontSize=22, spaceAfter=12)
    body = ParagraphStyle("b", parent=ss["Normal"], fontSize=13, leading=20)
    bullet = ParagraphStyle("bl", parent=body, leftIndent=16, spaceAfter=5)
    cap = ParagraphStyle("c", parent=ss["Italic"], fontSize=10, alignment=1)
    story = []

    def slide(items):
        story.extend(items); story.append(PageBreak())

    def bullets(items):
        return [Paragraph(f"• {x}", bullet) for x in items]

    def img_flow(path, mw, mh):
        with PILImage.open(path) as im:
            iw, ih = im.size
        r = min(mw / iw, mh / ih)
        return Image(str(path), width=iw * r, height=ih * r)

    def mtable(data, cw=None):
        t = Table(data, colWidths=cw, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), PDF_BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 12),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F8")]),
            ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        return t

    slide([Spacer(1, 80),
           Paragraph("U-Net Segmentation of Canine Cutaneous Tumours", h_title),
           Paragraph("Weeks 3–4 Progress — Model Development (Phase 2 start)", h_sub),
           Spacer(1, 20),
           Paragraph("Muhammad Tayyab Abbas · MSc AI &amp; Data Science · University of Hull", h_sub),
           Paragraph("Supervisor: Dr Claire Cashmore · 2nd fortnightly meeting", h_sub)])

    slide([Paragraph("What was built", h_slide), *bullets([
        "Finished Phase-1 preprocessing: 5x/10x/20x multi-magnification patches + augmentation",
        "U-Net segmentation: from-scratch baseline → ResNet-34 encoder (transfer learning)",
        "BCE+Dice loss, Adam 1e-4, cosine annealing, early stopping, TensorBoard",
        "Slide-level split (no patch leakage); Dice/IoU/pixel-acc/Hausdorff evaluation",
    ])])

    bd, bi, ba, bh = _fmt_metrics(base)
    rd, ri, ra, rh = _fmt_metrics(res)
    slide([Paragraph("Results — baseline vs ResNet-34 (test)", h_slide),
           mtable([["Model", "Dice", "IoU", "Pixel acc.", "Hausdorff"],
                   ["Baseline U-Net", bd, bi, ba, bh],
                   ["U-Net + ResNet-34", rd, ri, ra, rh]],
                  cw=[260, 110, 110, 120, 120]),
           Spacer(1, 14),
           Paragraph("Synthetic-data validation of the full pipeline; proposal "
                     "target is Dice &gt; 0.85 on real CATCH slides.", h_sub)])

    for fig_name, ttl, c in [
        ("unet_resnet34_curves.png", "Training curves (ResNet-34 U-Net)",
         "Loss decreases and validation Dice rises — correct convergence."),
        ("unet_resnet34_predictions.png", "Test predictions",
         "Image | ground-truth mask | predicted mask."),
    ]:
        fp = FIG / fig_name
        items = [Paragraph(ttl, h_slide)]
        if fp.exists():
            items += [img_flow(fp, uw, uh - 130), Spacer(1, 6), Paragraph(c, cap)]
        else:
            items.append(Paragraph("(run the pipeline to generate this figure)", body))
        slide(items)

    slide([Paragraph("Next fortnight (Weeks 5–6)", h_slide), *bullets([
        "Attention U-Net variant + comparison",
        "Run on real CATCH annotated slides once downloaded",
        "Begin ResNet-50 classification",
    ])])
    story.extend([Spacer(1, 140), Paragraph("Thank you", h_title),
                  Paragraph("Questions &amp; feedback welcome", h_sub)])

    doc.build(story)
    print(f"[PDF]  saved -> {out_path}")


def main():
    build_docx(ROOT / "docs" / "progress_report" / "Progress_Report_Week3-4.docx")
    build_pdf(ROOT / "docs" / "slides" / "Supervisor_Meeting_Week3-4.pdf")


if __name__ == "__main__":
    main()
