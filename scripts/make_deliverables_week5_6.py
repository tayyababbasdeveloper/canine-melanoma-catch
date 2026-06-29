"""Generate Week 5-6 supervisor deliverables.

  1. Progress_Report_Week5-6.docx
  2. Supervisor_Meeting_Week5-6.pdf  (Attention U-Net 3-way + ResNet-50 confusion)

Reads the JSON reports / figures produced by run_week5_6_pipeline.py. Run AFTER
the pipeline so the numbers and images exist.
"""
from __future__ import annotations

import sys
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "outputs" / "figures"
LOGS = ROOT / "outputs" / "logs"


def _load(name):
    p = LOGS / name
    return json.loads(p.read_text()) if p.exists() else None


def _seg_rows():
    cmp = _load("week5_6_seg_comparison.json") or {}
    nice = {"unet_baseline": "Baseline U-Net", "unet_resnet34": "U-Net + ResNet-34",
            "unet_attention": "Attention U-Net"}
    rows = [["Model", "Params", "Dice", "IoU"]]
    for k, label in nice.items():
        if k in cmp:
            v = cmp[k]
            rows.append([label, f"{v.get('params_million','—')}M",
                         f"{v['test_dice']:.3f}", f"{v['test_iou']:.3f}"])
    return rows


def _cls_metrics():
    r = _load("cls_resnet50_test_report.json")
    if not r:
        return None
    m = r["test_metrics"]
    return r, m


def build_docx(out_path: Path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HULL = RGBColor(0x1F, 0x3A, 0x6E)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def heading(t, lvl=1):
        h = doc.add_heading(t, level=lvl)
        for r in h.runs:
            r.font.color.rgb = HULL
        return h

    def table(rows, header=True):
        t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
                if header and i == 0:
                    for p in cells[j].paragraphs:
                        for rr in p.runs:
                            rr.font.bold = True
        return t

    title = doc.add_heading("Progress Report — Weeks 5–6 (Attention U-Net + "
                            "ResNet-50 Classification)", level=0)
    for r in title.runs:
        r.font.color.rgb = HULL
    for k, v in [("Project", "Classification of Malignant Melanoma in Canines (CATCH)"),
                 ("Project Code", "DAIM2025A_088"), ("Student", "Muhammad Tayyab Abbas"),
                 ("Supervisor", "Dr Claire Cashmore"),
                 ("Reporting period", "Weeks 5–6 (23 June – 6 July 2026)"),
                 ("Meeting", "Third fortnightly supervisor meeting")]:
        doc.add_paragraph(f"{k}: {v}")

    heading("1. Summary", 1)
    doc.add_paragraph(
        "Phase 2 continued on schedule. The Attention U-Net (Oktay et al., 2018) "
        "was implemented and trained, giving a three-way segmentation comparison "
        "(baseline / ResNet-34 / Attention). In parallel, the tumour-subtype "
        "classification track began: a ResNet-50 classifier (ImageNet transfer "
        "learning, progressive unfreezing, label-smoothing) was trained and "
        "evaluated with accuracy, macro-F1, AUC-ROC and a confusion matrix. All "
        "validated on synthetic data; the code runs unchanged on real CATCH slides.")

    heading("2. Planned vs. actual (Gantt)", 1)
    table([["Gantt task", "Planned", "Status"],
           ["Attention U-Net Implementation", "W5–6", "Complete (3-way comparison)"],
           ["ResNet-50 Classification", "W6–7", "Started (W6) — trained & evaluated"]])

    heading("3. Attention U-Net — segmentation (3-way)", 1)
    table(_seg_rows())
    doc.add_paragraph(
        "The attention gates re-weight the encoder skip features using the decoder "
        "signal, suppressing background and focusing on tumour regions. Built on the "
        "same backbone as the baseline, so the comparison isolates the attention "
        "effect.")
    fp = FIG / "unet_attention_predictions.png"
    if fp.exists():
        cap = doc.add_paragraph("Figure: Attention U-Net test predictions "
                                "(image | ground truth | prediction).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(fp), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading("4. ResNet-50 — tumour-subtype classification", 1)
    cm = _cls_metrics()
    if cm:
        r, m = cm
        table([["Metric", "Value"],
               ["Classes", ", ".join(r["classes"])],
               ["Accuracy", f"{m['accuracy']:.3f}"],
               ["Macro F1", f"{m['f1']:.3f}"],
               ["Macro precision", f"{m['precision']:.3f}"],
               ["Macro recall", f"{m['recall']:.3f}"],
               ["Macro AUC-ROC", "n/a" if m['auc_roc'] != m['auc_roc'] else f"{m['auc_roc']:.3f}"]])
        cfp = FIG / "cls_resnet50_confusion.png"
        if cfp.exists():
            cap = doc.add_paragraph("Figure: ResNet-50 confusion matrix (test).")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(cfp), width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("(Run the pipeline to populate classification results.)")

    doc.add_paragraph(
        "Trained with cross-entropy + label smoothing (0.1), AdamW, and progressive "
        "unfreezing (head → deeper layers). Note: synthetic-data validation; absolute "
        "values will change on real CATCH slides.")

    heading("5. Next fortnight (Weeks 7–8)", 1)
    for b in ["Train EfficientNet-B3 classifier and compare with ResNet-50.",
              "Run on the real CATCH slides once downloaded.",
              "Begin evaluation phase: 5-fold cross-validation, McNemar's test."]:
        doc.add_paragraph(b, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[DOCX] saved -> {out_path}")


def build_pdf(out_path: Path):
    from PIL import Image as PILImage
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, Image, PageBreak)
    BLUE = colors.HexColor("#1F3A6E")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    page = landscape(A4)
    doc = SimpleDocTemplate(str(out_path), pagesize=page, leftMargin=42, rightMargin=42,
                            topMargin=40, bottomMargin=40)
    uw, uh = page[0] - 84, page[1] - 80
    ss = getSampleStyleSheet()
    h_title = ParagraphStyle("t", parent=ss["Title"], textColor=BLUE, fontSize=27)
    h_sub = ParagraphStyle("s", parent=ss["Normal"], fontSize=14)
    h_slide = ParagraphStyle("sh", parent=ss["Heading1"], textColor=BLUE, fontSize=22, spaceAfter=12)
    body = ParagraphStyle("b", parent=ss["Normal"], fontSize=13, leading=20)
    bullet = ParagraphStyle("bl", parent=body, leftIndent=16, spaceAfter=5)
    cap = ParagraphStyle("c", parent=ss["Italic"], fontSize=10, alignment=1)
    story = []

    def slide(items): story.extend(items); story.append(PageBreak())
    def bullets(xs): return [Paragraph(f"• {x}", bullet) for x in xs]

    def img_flow(path, mw, mh):
        with PILImage.open(path) as im:
            iw, ih = im.size
        r = min(mw / iw, mh / ih)
        return Image(str(path), width=iw * r, height=ih * r)

    def mtable(data, cw=None):
        t = Table(data, colWidths=cw, hAlign="LEFT")
        t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 12), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F8")]),
            ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
        return t

    slide([Spacer(1, 80),
           Paragraph("Attention U-Net &amp; ResNet-50 Classification", h_title),
           Paragraph("Weeks 5–6 Progress — Model Development", h_sub), Spacer(1, 20),
           Paragraph("Muhammad Tayyab Abbas · MSc AI &amp; Data Science · University of Hull", h_sub),
           Paragraph("Supervisor: Dr Claire Cashmore · 3rd fortnightly meeting", h_sub)])

    slide([Paragraph("Segmentation — 3-way comparison", h_slide),
           mtable(_seg_rows(), cw=[230, 90, 110, 110]), Spacer(1, 12),
           Paragraph("Attention gates focus the skip connections on tumour regions; "
                     "built on the baseline backbone to isolate the attention effect.", h_sub)])

    fp = FIG / "unet_attention_predictions.png"
    items = [Paragraph("Attention U-Net — test predictions", h_slide)]
    if fp.exists():
        items += [img_flow(fp, uw, uh - 130), Spacer(1, 6),
                  Paragraph("Image | ground-truth mask | predicted mask.", cap)]
    slide(items)

    cm = _cls_metrics()
    if cm:
        r, m = cm
        slide([Paragraph("Classification — ResNet-50 (test)", h_slide),
               mtable([["Metric", "Value"], ["Accuracy", f"{m['accuracy']:.3f}"],
                       ["Macro F1", f"{m['f1']:.3f}"], ["Macro precision", f"{m['precision']:.3f}"],
                       ["Macro recall", f"{m['recall']:.3f}"],
                       ["AUC-ROC", "n/a" if m['auc_roc'] != m['auc_roc'] else f"{m['auc_roc']:.3f}"]],
                      cw=[260, 160]), Spacer(1, 10),
               Paragraph("Classes: " + ", ".join(r["classes"]), h_sub)])
        cfp = FIG / "cls_resnet50_confusion.png"
        items = [Paragraph("ResNet-50 — confusion matrix", h_slide)]
        if cfp.exists():
            items += [img_flow(cfp, uw, uh - 130), Spacer(1, 6),
                      Paragraph("Diagonal = correct subtype predictions.", cap)]
        slide(items)

    slide([Paragraph("Next fortnight (Weeks 7–8)", h_slide), *bullets([
        "EfficientNet-B3 classifier + comparison with ResNet-50",
        "Run on real CATCH slides once downloaded",
        "Evaluation: 5-fold cross-validation, McNemar's test"])])
    story.extend([Spacer(1, 140), Paragraph("Thank you", h_title),
                  Paragraph("Questions &amp; feedback welcome", h_sub)])
    doc.build(story)
    print(f"[PDF]  saved -> {out_path}")


def main():
    build_docx(ROOT / "docs" / "progress_report" / "Progress_Report_Week5-6.docx")
    build_pdf(ROOT / "docs" / "slides" / "Supervisor_Meeting_Week5-6.pdf")


if __name__ == "__main__":
    main()
