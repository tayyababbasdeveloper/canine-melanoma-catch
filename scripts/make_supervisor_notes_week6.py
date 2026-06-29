"""Generate the Week 6 supervisor-meeting notes as Word (.docx) and PDF.

Content mirrors docs/Supervisor_Notes_Week6.md. The segmentation/classification
numbers are pulled live from the JSON reports so the documents never drift from
the actual outputs.

    python scripts/make_supervisor_notes_week6.py
"""
from __future__ import annotations

import sys
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
ROOT = Path(__file__).resolve().parents[1]
LOGS = ROOT / "outputs" / "logs"
HULL_HEX = "#1F3A6E"


def _load(name):
    p = LOGS / name
    return json.loads(p.read_text()) if p.exists() else None


def _seg_rows():
    cmp = _load("week5_6_seg_comparison.json") or {}
    nice = {"unet_baseline": "Baseline U-Net", "unet_resnet34": "U-Net + ResNet-34",
            "unet_attention": "Attention U-Net"}
    rows = [["Model", "Params", "Dice", "IoU", "Hausdorff"]]
    for k, label in nice.items():
        if k in cmp:
            v = cmp[k]
            rows.append([label, f"{v.get('params_million','-')}M",
                         f"{v['test_dice']:.3f}", f"{v['test_iou']:.3f}",
                         f"{v.get('test_hausdorff', float('nan')):.1f}"])
    return rows


def _cls_line():
    r = _load("cls_resnet50_test_report.json")
    if not r:
        return "Classification report not found."
    m = r["test_metrics"]
    return (f"ResNet-50 (demo, {len(r['classes'])} classes): "
            f"accuracy={m['accuracy']:.3f}, macro-F1={m['f1']:.3f}, "
            f"AUC={m['auc_roc']:.3f} — perfect score is a synthetic-data artefact.")


# ----------------------------------------------------------------------- DOCX
def build_docx(out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    HULL = RGBColor(0x1F, 0x3A, 0x6E)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def heading(t, lvl=1):
        h = doc.add_heading(t, level=lvl)
        for r in h.runs:
            r.font.color.rgb = HULL
        return h

    title = doc.add_heading("Supervisor Meeting Notes — Week 6", level=0)
    for r in title.runs:
        r.font.color.rgb = HULL
    for k, v in [("Project", "Classification of Malignant Melanoma in Canines (CATCH)"),
                 ("Project Code", "DAIM2025A_088"),
                 ("Student", "Muhammad Tayyab Abbas"),
                 ("Supervisor", "Dr Claire Cashmore"),
                 ("Meeting", "Third fortnightly supervisor meeting — 29 June 2026")]:
        doc.add_paragraph(f"{k}: {v}")

    heading("1. Done this fortnight (W5-6)")
    for b in ["Attention U-Net (Oktay 2018) + 3-way segmentation comparison.",
              "ResNet-50 tumour-subtype classifier (transfer learning).",
              "Made the pipeline genuinely real-CATCH-ready (.svs reading, "
              "COCO/SQLite annotation masks, 7 subtypes, slide-level split)."]:
        doc.add_paragraph(b, style="List Bullet")

    heading("2. Honest data status")
    doc.add_paragraph(
        "No real CATCH slides processed yet — results are on SYNTHETIC data and "
        "validate the code, not the biology. The real set is 350 .svs WSIs, 50 each "
        "of 7 subtypes, 12,424 polygon annotations (TCIA DOI 10.7937/TCIA.2M93-FX66). "
        "Blocker: license-gated, large (~hundreds of GB) download; needs OpenSlide.")

    heading("3. Three bugs found and fixed")
    for b in ["Label bug: subtype was guessed from a filename colour-cast -> now "
              "from the real subtype folder/annotation.",
              "Data leakage: patch split mixed same-slide patches across train/test "
              "-> now slide-level (verified zero overlap).",
              "Transfer learning: frozen-encoder BatchNorm kept updating and "
              "overwriting pretrained stats -> now held in eval mode."]:
        doc.add_paragraph(b, style="List Bullet")

    heading("4. Results (synthetic)")
    rows = _seg_rows()
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            if i == 0:
                for p in cells[j].paragraphs:
                    for rr in p.runs:
                        rr.font.bold = True
    doc.add_paragraph(_cls_line())

    heading("5. Decisions I need from you")
    for b in ["7-class subtype vs. binary melanoma-vs-rest as the primary target?",
              "Combined end-to-end metric vs. separate seg/cls evaluations?",
              "Data scope: a subset of subtypes to start, or wait for all 350?",
              "Single split now, 5-fold CV reserved for the final?"]:
        doc.add_paragraph(b, style="List Bullet")

    heading("6. Next fortnight (W7-8)")
    for b in ["EfficientNet-B3 + comparison with ResNet-50.",
              "First real-data run once OpenSlide + download are in place.",
              "Evaluation phase: 5-fold CV, McNemar's test, Grad-CAM."]:
        doc.add_paragraph(b, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[DOCX] saved -> {out_path}")


# ------------------------------------------------------------------------ PDF
def build_pdf(out_path: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, ListFlowable, ListItem)
    BLUE = colors.HexColor(HULL_HEX)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, leftMargin=48,
                            rightMargin=48, topMargin=48, bottomMargin=40)
    ss = getSampleStyleSheet()
    h0 = ParagraphStyle("h0", parent=ss["Title"], textColor=BLUE, fontSize=22)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], textColor=BLUE, fontSize=14,
                        spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("b", parent=ss["Normal"], fontSize=10.5, leading=15)
    meta = ParagraphStyle("m", parent=body, textColor=colors.HexColor("#444444"))
    story = []

    def bullets(xs):
        return ListFlowable([ListItem(Paragraph(x, body), leftIndent=10) for x in xs],
                            bulletType="bullet", start="•")

    story += [Paragraph("Supervisor Meeting Notes — Week 6", h0), Spacer(1, 6),
              Paragraph("Classification of Malignant Melanoma in Canines (CATCH) · "
                        "DAIM2025A_088", meta),
              Paragraph("Muhammad Tayyab Abbas · Supervisor: Dr Claire Cashmore · "
                        "29 June 2026", meta), Spacer(1, 10)]

    story += [Paragraph("1. Done this fortnight", h1),
              bullets(["Attention U-Net (Oktay 2018) + 3-way segmentation comparison.",
                       "ResNet-50 tumour-subtype classifier (transfer learning).",
                       "Pipeline made real-CATCH-ready: .svs reading, COCO/SQLite "
                       "annotation masks, 7 subtypes, slide-level split."])]

    story += [Paragraph("2. Honest data status", h1),
              Paragraph("No real CATCH slides processed yet — results are on "
                        "<b>synthetic</b> data and validate the code, not the biology. "
                        "Real set: 350 .svs WSIs, 50 each of 7 subtypes, 12,424 polygon "
                        "annotations (TCIA DOI 10.7937/TCIA.2M93-FX66). Blocker: "
                        "license-gated, ~hundreds-of-GB download; needs OpenSlide.", body)]

    story += [Paragraph("3. Three bugs found and fixed", h1),
              bullets(["<b>Label bug</b>: subtype guessed from a filename colour-cast "
                       "→ now from the real subtype folder/annotation.",
                       "<b>Data leakage</b>: patch split mixed same-slide patches "
                       "across train/test → now slide-level (verified zero overlap).",
                       "<b>Transfer learning</b>: frozen-encoder BatchNorm kept "
                       "updating → now held in eval mode."])]

    rows = _seg_rows()
    tbl = Table(rows, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F8")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story += [Paragraph("4. Results (synthetic — pipeline check, not biology)", h1),
              tbl, Spacer(1, 5), Paragraph(_cls_line(), body)]

    story += [Paragraph("5. Decisions I need from you", h1),
              bullets(["7-class subtype vs. binary melanoma-vs-rest as primary target?",
                       "Combined end-to-end metric vs. separate seg/cls evaluations?",
                       "Data scope: a subset of subtypes to start, or wait for all 350?",
                       "Single split now, 5-fold CV reserved for the final?"])]

    story += [Paragraph("6. Next fortnight (W7-8)", h1),
              bullets(["EfficientNet-B3 + comparison with ResNet-50.",
                       "First real-data run once OpenSlide + download are in place.",
                       "Evaluation phase: 5-fold CV, McNemar's test, Grad-CAM."])]

    doc.build(story)
    print(f"[PDF]  saved -> {out_path}")


def main():
    build_docx(ROOT / "docs" / "Supervisor_Notes_Week6.docx")
    build_pdf(ROOT / "docs" / "Supervisor_Notes_Week6.pdf")


if __name__ == "__main__":
    main()
